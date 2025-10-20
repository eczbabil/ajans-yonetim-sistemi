"""
Word Rapor Oluşturma Modülü
Müşteri bazlı detaylı raporlar oluşturur
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
from collections import Counter


def add_heading_custom(doc, text, level=1):
    """
    Özel başlık ekler
    """
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(37, 99, 235)  # Mavi renk
    return heading


def generate_musteri_raporu(db, musteri, start_date=None, end_date=None):
    """
    Müşteri için detaylı Word raporu oluşturur
    
    Args:
        db: Database session
        musteri: Musteri objesi
        start_date: Başlangıç tarihi
        end_date: Bitiş tarihi
    
    Returns:
        Document: Word dokümanı
    """
    try:
        from app import IsGunlugu, Teslimat, SosyalMedya, Revizyon
    except ImportError:
        # Alternatif import yolu
        from ...app import IsGunlugu, Teslimat, SosyalMedya, Revizyon
    
    # Tarih aralığı belirleme
    if not end_date:
        end_date = datetime.now().date()
    if not start_date:
        # Son 30 gün
        start_date = end_date - timedelta(days=30)
    
    # Word dokümanı oluştur
    doc = Document()
    
    # Başlık
    title = doc.add_heading(f'{musteri.ad} - Detaylı Ajans Raporu', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tarih bilgisi
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(
        f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}\n'
        f'Dönem: {start_date.strftime("%d.%m.%Y")} - {end_date.strftime("%d.%m.%Y")}'
    )
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Boşluk
    
    # ===== GENEL ÖZET =====
    add_heading_custom(doc, '📊 Genel Özet', level=1)
    
    # İş günlüğü verileri
    isler = db.session.query(IsGunlugu).filter(
        IsGunlugu.musteri_id == musteri.id,
        IsGunlugu.tarih >= start_date,
        IsGunlugu.tarih <= end_date
    ).all()
    
    toplam_dakika = sum([is_item.sure_dakika for is_item in isler])
    toplam_saat = round(toplam_dakika / 60, 1)
    is_sayisi = len(isler)
    
    # Benzersiz iş günü sayısı
    benzersiz_gunler = len(set([is_item.tarih for is_item in isler]))
    
    # Benzersiz sorumlu kişiler
    sorumlu_kisiler = set([is_item.sorumlu_kisi for is_item in isler if is_item.sorumlu_kisi])
    ekip_buyuklugu = len(sorumlu_kisiler)
    
    # Özet paragraf
    ozet = doc.add_paragraph()
    ozet_text = (
        f"Toplam ortalama {toplam_saat} saat hizmet verilmiş. "
        f"{musteri.ad} için {benzersiz_gunler} iş günü boyunca "
        f"{ekip_buyuklugu} kişilik ekibimiz toplam {toplam_saat} saat çalıştı."
    )
    ozet.add_run(ozet_text).bold = True
    
    doc.add_paragraph()
    
    # ===== GENEL GÖSTERGELER TABLOSU =====
    add_heading_custom(doc, '📈 Genel Göstergeler', level=2)
    
    # Teslimat verileri
    teslimatlar = db.session.query(Teslimat).filter(
        Teslimat.musteri_id == musteri.id,
        Teslimat.teslim_tarihi >= start_date,
        Teslimat.teslim_tarihi <= end_date
    ).all()
    
    # Teslimat türlerine göre say
    tasarim_sayisi = len([t for t in teslimatlar if 'tasarım' in (t.teslim_turu or '').lower() or 'tasarım' in (t.aktivite_turu or '').lower()])
    video_sayisi = len([t for t in teslimatlar if 'video' in (t.teslim_turu or '').lower() or 'video' in (t.aktivite_turu or '').lower()])
    
    # Revizyonlar
    revizyonlar = db.session.query(Revizyon).filter(
        Revizyon.musteri_id == musteri.id,
        Revizyon.tarih >= start_date,
        Revizyon.tarih <= end_date
    ).all()
    
    tasarim_revize = len([r for r in revizyonlar if any(t.id == r.teslimat_id and 'tasarım' in (t.teslim_turu or '').lower() for t in teslimatlar)])
    video_revize = len([r for r in revizyonlar if any(t.id == r.teslimat_id and 'video' in (t.teslim_turu or '').lower() for t in teslimatlar)])
    
    onaylanan = len([t for t in teslimatlar if t.durum == 'Onaylandı'])
    
    # Sosyal medya
    sosyal_medyalar = db.session.query(SosyalMedya).filter(
        SosyalMedya.musteri_id == musteri.id,
        SosyalMedya.tarih >= start_date,
        SosyalMedya.tarih <= end_date
    ).all()
    
    yayinlanan = len([s for s in sosyal_medyalar if s.durum == 'Yayınlandı'])
    
    # Tablo oluştur
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Başlıklar
    table.rows[0].cells[0].text = 'Gösterge'
    table.rows[0].cells[1].text = 'Sonuç'
    
    # Veriler
    table.rows[1].cells[0].text = 'Toplam medya dosyası'
    table.rows[1].cells[1].text = f"{len(teslimatlar)} ({tasarim_sayisi} tasarım + {video_sayisi} video)"
    
    table.rows[2].cells[0].text = 'Toplam revize'
    table.rows[2].cells[1].text = f"{len(revizyonlar)} ({tasarim_revize} tasarım + {video_revize} video)"
    
    table.rows[3].cells[0].text = 'Onaylanan işler'
    table.rows[3].cells[1].text = str(onaylanan)
    
    table.rows[4].cells[0].text = 'Yayınlanan işler'
    table.rows[4].cells[1].text = str(yayinlanan)
    
    table.rows[5].cells[0].text = 'Red edilen'
    table.rows[5].cells[1].text = "0"
    
    table.rows[6].cells[0].text = 'Proje süresi'
    table.rows[6].cells[1].text = f"{(end_date - start_date).days} gün ({start_date.strftime('%d %b')} – {end_date.strftime('%d %b %Y')})"
    
    doc.add_page_break()
    
    # ===== KİŞİ BAŞI İŞ DAĞILIMI =====
    add_heading_custom(doc, '👥 Kişi Başı İş Dağılımı (Top 5)', level=2)
    
    # Sorumlu kişilere göre iş sayısı ve süre
    kisi_istatistik = {}
    for is_item in isler:
        if is_item.sorumlu_kisi:
            if is_item.sorumlu_kisi not in kisi_istatistik:
                kisi_istatistik[is_item.sorumlu_kisi] = {'is_sayisi': 0, 'toplam_dakika': 0}
            kisi_istatistik[is_item.sorumlu_kisi]['is_sayisi'] += 1
            kisi_istatistik[is_item.sorumlu_kisi]['toplam_dakika'] += is_item.sure_dakika
    
    # Top 5 kişi
    top_kisiler = sorted(kisi_istatistik.items(), key=lambda x: x[1]['is_sayisi'], reverse=True)[:5]
    
    if top_kisiler:
        kisi_table = doc.add_table(rows=len(top_kisiler) + 1, cols=3)
        kisi_table.style = 'Light Grid Accent 1'
        
        # Başlıklar
        kisi_table.rows[0].cells[0].text = 'İsim'
        kisi_table.rows[0].cells[1].text = 'İş Sayısı'
        kisi_table.rows[0].cells[2].text = 'Toplam Saat'
        
        for i, (kisi, stats) in enumerate(top_kisiler, 1):
            kisi_table.rows[i].cells[0].text = kisi
            kisi_table.rows[i].cells[1].text = str(stats['is_sayisi'])
            kisi_table.rows[i].cells[2].text = f"{round(stats['toplam_dakika'] / 60, 1)} saat"
    else:
        doc.add_paragraph("Kişi bazlı iş verisi bulunmuyor.")
    
    doc.add_paragraph()
    
    # ===== TESLİMAT REVİZE & ONAY SÜRECİ =====
    add_heading_custom(doc, '📦 Teslimat Revize & Onay Süreci', level=2)
    
    if teslimatlar:
        # Tablo oluştur
        teslimat_table = doc.add_table(rows=len(teslimatlar) + 1, cols=6)
        teslimat_table.style = 'Light Grid Accent 1'
        
        # Başlıklar
        teslimat_table.rows[0].cells[0].text = 'Dosya Adı'
        teslimat_table.rows[0].cells[1].text = 'Gönderim Tarihi'
        teslimat_table.rows[0].cells[2].text = 'Gönderen'
        teslimat_table.rows[0].cells[3].text = 'Revize Durumu'
        teslimat_table.rows[0].cells[4].text = 'Onay'
        teslimat_table.rows[0].cells[5].text = 'Yayın/Paylaşım'
        
        for i, teslimat in enumerate(teslimatlar, 1):
            # Revizyon sayısını hesapla
            teslimat_revizyonlari = [r for r in revizyonlar if r.teslimat_id == teslimat.id]
            revize_durumu = f"{len(teslimat_revizyonlari)} Revizyon" if teslimat_revizyonlari else "Revize yok"
            
            # Yayın bilgisi - sosyal medyadan kontrol
            yayin_bilgisi = "-"
            for sosyal in sosyal_medyalar:
                if teslimat.baslik and teslimat.baslik.lower() in (sosyal.icerik_basligi or '').lower():
                    if sosyal.durum == 'Yayınlandı':
                        yayin_bilgisi = f"{sosyal.platform}'da yayınlandı"
                        break
            
            teslimat_table.rows[i].cells[0].text = teslimat.baslik or 'Başlıksız'
            teslimat_table.rows[i].cells[1].text = teslimat.teslim_tarihi.strftime('%d.%m.%Y %H:%M') if teslimat.teslim_tarihi else '-'
            teslimat_table.rows[i].cells[2].text = teslimat.sorumlu_kisi or 'Belirsiz'
            teslimat_table.rows[i].cells[3].text = revize_durumu
            teslimat_table.rows[i].cells[4].text = teslimat.durum or 'Bekliyor'
            teslimat_table.rows[i].cells[5].text = yayin_bilgisi
    else:
        doc.add_paragraph("Bu dönemde teslimat bulunmuyor.")
    
    doc.add_paragraph()
    
    # ===== REVİZYON DETAYLARI TABLOSU =====
    add_heading_custom(doc, '🔄 Revizyon Durumları', level=2)
    
    if revizyonlar:
        # Tablo oluştur
        revizyon_table = doc.add_table(rows=len(revizyonlar) + 1, cols=4)
        revizyon_table.style = 'Light Grid Accent 1'
        
        # Başlıklar
        revizyon_table.rows[0].cells[0].text = 'Tarih'
        revizyon_table.rows[0].cells[1].text = 'Dosya/Tasarım'
        revizyon_table.rows[0].cells[2].text = 'Revize Nedeni'
        revizyon_table.rows[0].cells[3].text = 'Durum'
        
        for i, revizyon in enumerate(revizyonlar, 1):
            # Teslimat bilgisini bul
            teslimat = next((t for t in teslimatlar if t.id == revizyon.teslimat_id), None)
            teslimat_adi = teslimat.baslik if teslimat else f"Teslimat #{revizyon.teslimat_id}"
            
            revizyon_table.rows[i].cells[0].text = revizyon.tarih.strftime('%d.%m.%Y')
            revizyon_table.rows[i].cells[1].text = teslimat_adi
            revizyon_table.rows[i].cells[2].text = revizyon.revize_konusu[:50] if revizyon.revize_konusu else '-'
            revizyon_table.rows[i].cells[3].text = revizyon.durum or 'Bekliyor'
    else:
        doc.add_paragraph("Bu dönemde revizyon bulunmuyor.")
    
    doc.add_paragraph()
    
    # ===== DURUM ÖZETİ TABLOSU =====
    add_heading_custom(doc, '📋 Durum Özeti', level=2)
    
    durum_table = doc.add_table(rows=5, cols=2)
    durum_table.style = 'Light Grid Accent 1'
    
    durum_table.rows[0].cells[0].text = 'Kategori'
    durum_table.rows[0].cells[1].text = 'Adet'
    
    durum_table.rows[1].cells[0].text = 'Toplam iş üretilen'
    durum_table.rows[1].cells[1].text = str(len(teslimatlar))
    
    durum_table.rows[2].cells[0].text = 'Revize alan'
    durum_table.rows[2].cells[1].text = str(len(revizyonlar))
    
    durum_table.rows[3].cells[0].text = 'Onaylanan'
    durum_table.rows[3].cells[1].text = str(onaylanan)
    
    durum_table.rows[4].cells[0].text = 'Yayınlanan'
    durum_table.rows[4].cells[1].text = str(yayinlanan)
    
    doc.add_paragraph()
    
    # ===== İŞ TİPİ DAĞILIMI =====
    add_heading_custom(doc, '🎨 İş Tipi Dağılımı', level=2)
    
    aktivite_turleri = [is_item.aktivite_turu for is_item in isler if is_item.aktivite_turu]
    aktivite_sayaci = Counter(aktivite_turleri)
    
    if aktivite_sayaci:
        # Tablo formatında göster
        is_tipi_table = doc.add_table(rows=len(aktivite_sayaci) + 1, cols=3)
        is_tipi_table.style = 'Light Grid Accent 1'
        
        # Başlıklar
        is_tipi_table.rows[0].cells[0].text = 'Aktivite Türü'
        is_tipi_table.rows[0].cells[1].text = 'İş Sayısı'
        is_tipi_table.rows[0].cells[2].text = 'Toplam Saat'
        
        # Aktivite türüne göre süre hesapla
        for i, (aktivite, sayi) in enumerate(aktivite_sayaci.most_common(), 1):
            toplam_dakika_aktivite = sum([is_item.sure_dakika for is_item in isler if is_item.aktivite_turu == aktivite])
            is_tipi_table.rows[i].cells[0].text = aktivite
            is_tipi_table.rows[i].cells[1].text = str(sayi)
            is_tipi_table.rows[i].cells[2].text = f"{round(toplam_dakika_aktivite / 60, 1)} saat"
    else:
        doc.add_paragraph("İş tipi verisi bulunmuyor.")
    
    doc.add_paragraph()
    
    # ===== TESLİM TAKVİMİ =====
    add_heading_custom(doc, '📅 Teslim Takvimi - Teslimat Çıktıları', level=2)
    
    if teslimatlar:
        # Tarihe göre grupla
        from itertools import groupby
        teslimatlar_sorted = sorted(teslimatlar, key=lambda t: t.teslim_tarihi if t.teslim_tarihi else datetime(1900, 1, 1).date())
        
        teslim_table = doc.add_table(rows=len(teslimatlar) + 1, cols=4)
        teslim_table.style = 'Light Grid Accent 1'
        
        # Başlıklar
        teslim_table.rows[0].cells[0].text = 'Teslim Tarihi'
        teslim_table.rows[0].cells[1].text = 'Tür'
        teslim_table.rows[0].cells[2].text = 'Dosyalar'
        teslim_table.rows[0].cells[3].text = 'Durum'
        
        for i, teslimat in enumerate(teslimatlar_sorted, 1):
            teslim_table.rows[i].cells[0].text = teslimat.teslim_tarihi.strftime('%d.%m.%Y') if teslimat.teslim_tarihi else '-'
            teslim_table.rows[i].cells[1].text = teslimat.teslim_turu or teslimat.aktivite_turu or '-'
            teslim_table.rows[i].cells[2].text = teslimat.baslik or '-'
            teslim_table.rows[i].cells[3].text = teslimat.durum or 'Bekliyor'
    else:
        doc.add_paragraph("Teslimat verisi bulunmuyor.")
    
    doc.add_paragraph()
    
    # ===== KATEGORİ BAZLI İÇERİK LİSTESİ =====
    add_heading_custom(doc, '🎨 Kategori Bazlı İçerik Listesi', level=2)
    
    # Teslimat türlerine göre grupla
    kategori_grup = {}
    for teslimat in teslimatlar:
        kategori = teslimat.teslim_turu or teslimat.aktivite_turu or 'Diğer'
        if kategori not in kategori_grup:
            kategori_grup[kategori] = []
        kategori_grup[kategori].append(teslimat)
    
    if kategori_grup:
        for kategori, items in sorted(kategori_grup.items()):
            doc.add_heading(f"{kategori} ({len(items)} adet)", level=3)
            for teslimat in items[:5]:  # İlk 5'ini göster
                doc.add_paragraph(
                    f"• {teslimat.baslik} - {teslimat.teslim_tarihi.strftime('%d.%m.%Y') if teslimat.teslim_tarihi else '-'} - {teslimat.sorumlu_kisi or '-'}",
                    style='List Bullet'
                )
            if len(items) > 5:
                doc.add_paragraph(f"  ... ve {len(items) - 5} adet daha")
    else:
        doc.add_paragraph("Kategori verisi bulunmuyor.")
    
    doc.add_paragraph()
    
    # ===== YAYINLANAN İÇERİKLER =====
    yayinlanan_icerikler = [s for s in sosyal_medyalar if s.durum == 'Yayınlandı']
    
    if yayinlanan_icerikler:
        add_heading_custom(doc, '🌐 Yayınlanan Grafik İçerikler', level=2)
        
        # Tablo oluştur
        yayin_table = doc.add_table(rows=len(yayinlanan_icerikler) + 1, cols=5)
        yayin_table.style = 'Light Grid Accent 1'
        
        # Başlıklar
        yayin_table.rows[0].cells[0].text = 'Tarih'
        yayin_table.rows[0].cells[1].text = 'İçerik'
        yayin_table.rows[0].cells[2].text = 'Platform'
        yayin_table.rows[0].cells[3].text = 'Görüntülenme'
        yayin_table.rows[0].cells[4].text = 'Beğeni'
        
        for i, sosyal in enumerate(yayinlanan_icerikler, 1):
            yayin_table.rows[i].cells[0].text = sosyal.tarih.strftime('%d.%m.%Y')
            yayin_table.rows[i].cells[1].text = f"{sosyal.icerik_basligi} ({sosyal.gonderi_turu})"
            yayin_table.rows[i].cells[2].text = sosyal.platform
            yayin_table.rows[i].cells[3].text = str(sosyal.goruntulenme or 0)
            yayin_table.rows[i].cells[4].text = str(sosyal.begeni or 0)
        
        doc.add_paragraph()
    
    # ===== SONUÇ =====
    doc.add_page_break()
    add_heading_custom(doc, '✅ Sonuç ve Özet', level=2)
    
    sonuc = doc.add_paragraph()
    sonuc.add_run(f"Ajans, {(end_date - start_date).days} gün içinde {musteri.ad} için:\n\n").bold = True
    
    # Teslimat özeti
    sonuc.add_run(f"• {len(teslimatlar)} teslimat üretmiş")
    if tasarim_sayisi > 0 or video_sayisi > 0:
        sonuc.add_run(f" ({tasarim_sayisi} tasarım + {video_sayisi} video)")
    sonuc.add_run("\n")
    
    # Revizyon özeti
    sonuc.add_run(f"• Bunların {len(revizyonlar)}'i revize edilmiş")
    if tasarim_revize > 0 or video_revize > 0:
        sonuc.add_run(f" ({tasarim_revize} tasarım + {video_revize} video)")
    sonuc.add_run("\n")
    
    # Onay ve yayın
    sonuc.add_run(f"• {onaylanan} tanesi onaylanmış\n")
    sonuc.add_run(f"• {len(yayinlanan_icerikler)} tanesi dijital yayına gitmiştir\n")
    
    # Çalışma saati
    sonuc.add_run(f"• Toplam {toplam_saat} saat hizmet verilmiştir\n")
    sonuc.add_run(f"• {benzersiz_gunler} iş günü boyunca çalışılmıştır\n")
    
    # Ekip bilgisi
    if ekip_buyuklugu > 0:
        sonuc.add_run(f"• {ekip_buyuklugu} kişilik ekip çalışmıştır\n")
    
    # Aktivite yoğunluğu
    if aktivite_sayaci:
        en_yogun_aktivite = aktivite_sayaci.most_common(1)[0]
        sonuc.add_run(f"\nAktivite yoğunluğu: {en_yogun_aktivite[0]} ({en_yogun_aktivite[1]} iş)")
    
    return doc


def save_musteri_raporu(doc, musteri, output_path=None):
    """
    Raporu dosyaya kaydeder
    
    Args:
        doc: Document objesi
        musteri: Musteri objesi
        output_path: Çıktı dosya yolu (opsiyonel)
    
    Returns:
        str: Dosya yolu
    """
    if not output_path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{musteri.ad.replace(' ', '_')}_{timestamp}_rapor.docx"
        output_path = f"uploads/{filename}"
    
    doc.save(output_path)
    return output_path

